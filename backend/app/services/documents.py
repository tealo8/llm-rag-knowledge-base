from __future__ import annotations

import hashlib
import io
import json
import re
import sqlite3
import unicodedata
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException
from openpyxl import load_workbook
from pypdf import PdfReader

from ..config import get_settings
from ..db import db_session, utc_now
from ..security import CurrentUser, require_knowledge_base_permission
from .audit import write_audit
from .embeddings import EmbeddingService, fts_text, search_terms
from .settings import get_org_settings
from .vector_store import VectorRecord, get_vector_store


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".xlsx"}


@dataclass(frozen=True)
class ParsedPage:
    page_number: int | None
    text: str


@dataclass(frozen=True)
class TextChunk:
    page_number: int | None
    paragraph_number: int | None
    text: str


def safe_filename(filename: str) -> str:
    cleaned = Path(filename).name.strip().replace("\x00", "")
    if not cleaned:
        raise HTTPException(status_code=400, detail="文件名无效")
    suffix = Path(cleaned).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail="仅支持 PDF、DOCX、XLSX、Markdown 和 TXT")
    return cleaned


def validate_file_signature(filename: str, content: bytes) -> None:
    """Validate a small set of magic bytes before handing data to parsers.

    Extension and client MIME are untrusted. This is intentionally a lightweight
    gate; production deployments should additionally scan uploads in a sandbox.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" and not content.lstrip().startswith(b"%PDF-"):
        # Keep the established parser-facing error for corrupt PDFs.
        raise HTTPException(status_code=422, detail="PDF 解析失败")
    if suffix in {".docx", ".xlsx"} and not content.startswith(b"PK"):
        label = suffix[1:].upper()
        raise HTTPException(status_code=422, detail=f"{label} 文件内容与扩展名不匹配")
    if suffix in {".txt", ".md"} and b"\x00" in content:
        raise HTTPException(status_code=422, detail="文本文件内容无效，疑似二进制或可执行文件")


def clean_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\u200b", "").replace("\ufeff", "")
    normalized = "".join(
        char for char in normalized if char in "\n\t" or unicodedata.category(char) != "Cc"
    )
    normalized = re.sub(r"[ \t]+", " ", normalized.replace("\r\n", "\n"))
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def parse_document(filename: str, content: bytes) -> list[ParsedPage]:
    suffix = Path(filename).suffix.lower()
    safe_filename(filename)
    validate_file_signature(filename, content)

    if suffix == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [
                ParsedPage(index + 1, clean_text(page.extract_text() or ""))
                for index, page in enumerate(reader.pages)
            ]
        except Exception as exc:
            raise HTTPException(status_code=422, detail="PDF 解析失败") from exc
        if pages and not any(page.text for page in pages):
            raise HTTPException(
                status_code=422,
                detail="PDF 未提取到文本，可能是扫描图片 PDF；当前未安装 OCR 组件",
            )
    elif suffix == ".docx":
        try:
            document = DocxDocument(io.BytesIO(content))
            blocks = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            for table_index, table in enumerate(document.tables, start=1):
                blocks.append(f"[表格 {table_index}]")
                blocks.extend(
                    " | ".join(clean_text(cell.text) for cell in row.cells)
                    for row in table.rows
                )
            pages = [ParsedPage(None, clean_text("\n".join(blocks)))]
        except Exception as exc:
            raise HTTPException(status_code=422, detail="DOCX 解析失败") from exc
    elif suffix == ".xlsx":
        try:
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            pages = []
            for sheet_number, sheet in enumerate(workbook.worksheets, start=1):
                rows = [f"[工作表] {sheet.title}"]
                for row in sheet.iter_rows(values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(value.strip() for value in values):
                        rows.append(" | ".join(values))
                text = clean_text("\n".join(rows))
                if text:
                    pages.append(ParsedPage(sheet_number, text))
            workbook.close()
        except Exception as exc:
            raise HTTPException(status_code=422, detail="XLSX 解析失败") from exc
    else:
        decoded = None
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                decoded = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if decoded is None:
            raise HTTPException(status_code=422, detail="文本文件编码无法识别")
        pages = [ParsedPage(None, clean_text(decoded))]

    pages = [page for page in pages if page.text]
    if not pages:
        raise HTTPException(status_code=422, detail="文档中没有可索引的文本")
    return pages


def _fixed_split(text: str, target: int, overlap: int) -> list[str]:
    if len(text) <= target:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        hard_end = min(start + target, len(text))
        end = hard_end
        if hard_end < len(text):
            search_start = min(start + target // 2, hard_end)
            candidates = [
                text.rfind(boundary, search_start, hard_end)
                for boundary in ("\n\n", "。", "！", "？", ". ", "; ")
            ]
            best = max(candidates)
            if best > search_start:
                end = best + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def _semantic_split(text: str, target: int, overlap: int) -> list[str]:
    units = [
        unit.strip()
        for unit in re.split(r"(?<=[。！？!?])\s*|\n+", text)
        if unit.strip()
    ]
    chunks: list[str] = []
    current: list[str] = []
    current_length = 0
    for unit in units:
        if current and current_length + len(unit) > target:
            chunks.append("\n".join(current))
            retained: list[str] = []
            retained_length = 0
            for previous in reversed(current):
                if retained_length + len(previous) > overlap:
                    break
                retained.insert(0, previous)
                retained_length += len(previous)
            current = retained
            current_length = retained_length
        current.append(unit)
        current_length += len(unit)
    if current:
        chunks.append("\n".join(current))
    return chunks


def chunk_pages(
    pages: list[ParsedPage],
    *,
    strategy: str,
    target: int,
    overlap: int,
) -> list[TextChunk]:
    if overlap >= target:
        raise HTTPException(status_code=422, detail="切块重叠度必须小于切块大小")
    chunks: list[TextChunk] = []
    paragraph = 0
    splitter = _semantic_split if strategy == "semantic" else _fixed_split
    for page in pages:
        for text in splitter(page.text, target, overlap):
            paragraph += 1
            chunks.append(TextChunk(page.page_number, paragraph, text))
    return chunks


def _validate_ids(
    connection: sqlite3.Connection,
    *,
    table: str,
    ids: list[str],
    org_id: str,
    label: str,
) -> list[str]:
    unique_ids = list(dict.fromkeys(ids))
    if not unique_ids:
        return []
    placeholders = ",".join("?" for _ in unique_ids)
    rows = connection.execute(
        f"SELECT id FROM {table} WHERE org_id = ? AND id IN ({placeholders})",
        [org_id, *unique_ids],
    ).fetchall()
    if len(rows) != len(unique_ids):
        raise HTTPException(status_code=403, detail=f"包含不属于当前组织的{label}")
    return unique_ids


def validate_document_acl(
    connection: sqlite3.Connection,
    org_id: str,
    visibility: str,
    group_ids: list[str],
    user_ids: list[str],
) -> tuple[list[str], list[str]]:
    groups = _validate_ids(
        connection, table="groups", ids=group_ids, org_id=org_id, label="用户组"
    )
    users = _validate_ids(
        connection, table="users", ids=user_ids, org_id=org_id, label="用户"
    )
    if visibility == "restricted" and not groups and not users:
        raise HTTPException(status_code=422, detail="受限文档至少需要选择用户组或指定用户")
    if visibility != "restricted":
        return [], []
    return groups, users


def validate_groups(
    connection: sqlite3.Connection,
    org_id: str,
    visibility: str,
    group_ids: list[str],
) -> list[str]:
    groups, _ = validate_document_acl(connection, org_id, visibility, group_ids, [])
    return groups


def resolve_knowledge_base(
    connection: sqlite3.Connection,
    user: CurrentUser,
    knowledge_base_id: str | None,
    minimum: str,
) -> str:
    if knowledge_base_id:
        require_knowledge_base_permission(connection, user, knowledge_base_id, minimum)
        return knowledge_base_id
    if user.role == "admin":
        row = connection.execute(
            "SELECT id FROM knowledge_bases WHERE org_id = ? AND status = 'active' ORDER BY created_at LIMIT 1",
            (user.org_id,),
        ).fetchone()
    else:
        row = connection.execute(
            """
            SELECT kb.id FROM knowledge_bases kb
            JOIN knowledge_base_access kba ON kba.knowledge_base_id = kb.id
            WHERE kb.org_id = ? AND kb.status = 'active' AND kba.user_id = ?
            ORDER BY kb.created_at LIMIT 1
            """,
            (user.org_id, user.id),
        ).fetchone()
    if row is None:
        raise HTTPException(status_code=404, detail="没有可用的知识库空间")
    require_knowledge_base_permission(connection, user, row["id"], minimum)
    return str(row["id"])


async def _index_document(document_id: str, content: bytes) -> int:
    with db_session() as connection:
        document = connection.execute(
            "SELECT * FROM documents WHERE id = ?", (document_id,)
        ).fetchone()
        if document is None:
            raise HTTPException(status_code=404, detail="文档不存在")
        connection.execute(
            "UPDATE documents SET status = 'processing', processing_stage = 'parsing', error = NULL WHERE id = ?",
            (document_id,),
        )
        previous_ids = [
            row["id"]
            for row in connection.execute(
                "SELECT id FROM chunks WHERE document_id = ?", (document_id,)
            ).fetchall()
        ]
    org_settings = get_org_settings(document["org_id"])
    strategy = document["chunk_strategy"] or org_settings["chunk_strategy"]
    pages = parse_document(document["filename"], content)
    chunks = chunk_pages(
        pages,
        strategy=strategy,
        target=int(org_settings["chunk_size"]),
        overlap=int(org_settings["chunk_overlap"]),
    )
    embedding_service = EmbeddingService()
    with db_session() as connection:
        connection.execute(
            "UPDATE documents SET processing_stage = 'embedding' WHERE id = ?",
            (document_id,),
        )
    embeddings = await embedding_service.embed_batched([chunk.text for chunk in chunks])
    chunk_ids = [str(uuid.uuid4()) for _ in chunks]
    vector_store = get_vector_store()
    vector_store.upsert(
        document["org_id"],
        document["knowledge_base_id"],
        [
            VectorRecord(chunk_id=chunk_id, document_id=document_id, vector=embedding)
            for chunk_id, embedding in zip(chunk_ids, embeddings, strict=True)
        ],
    )
    vector_store.delete_chunks(
        document["org_id"], document["knowledge_base_id"], previous_ids
    )
    with db_session() as connection:
        previous = connection.execute(
            "SELECT id FROM chunks WHERE document_id = ?", (document_id,)
        ).fetchall()
        connection.executemany(
            "DELETE FROM chunks_fts WHERE chunk_id = ?", [(row["id"],) for row in previous]
        )
        connection.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
        for position, (chunk_id, chunk, embedding) in enumerate(
            zip(chunk_ids, chunks, embeddings, strict=True)
        ):
            tokenized = fts_text(chunk.text)
            connection.execute(
                """
                INSERT INTO chunks
                    (id, document_id, org_id, knowledge_base_id, position, page_number,
                     paragraph_number, text, search_text, embedding_json, embedding_model,
                     embedding_dimensions, token_count, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    chunk_id,
                    document_id,
                    document["org_id"],
                    document["knowledge_base_id"],
                    position,
                    chunk.page_number,
                    chunk.paragraph_number,
                    chunk.text,
                    tokenized,
                    json.dumps(embedding),
                    embedding_service.backend_name,
                    len(embedding),
                    len(search_terms(chunk.text)),
                    utc_now(),
                ),
            )
            connection.execute(
                "INSERT INTO chunks_fts (chunk_id, title, search_text) VALUES (?, ?, ?)",
                (chunk_id, fts_text(document["title"]), tokenized),
            )
        connection.execute(
            "UPDATE documents SET status = 'indexed', processing_stage = 'indexed', chunk_count = ?, error = NULL WHERE id = ?",
            (len(chunks), document_id),
        )
    return len(chunks)


async def ingest_document(
    *,
    user: CurrentUser,
    filename: str,
    content_type: str,
    content: bytes,
    title: str,
    visibility: str,
    group_ids: list[str],
    user_ids: list[str] | None = None,
    knowledge_base_id: str | None = None,
    tags: list[str] | None = None,
    chunk_strategy: str | None = None,
    version_of: str | None = None,
    is_seed: bool = False,
) -> str:
    settings = get_settings()
    clean_name = safe_filename(filename)
    if len(content) > settings.max_chunked_upload_bytes:
        raise HTTPException(status_code=413, detail="文件超过服务端可配置上限")
    if not content:
        raise HTTPException(status_code=422, detail="文件内容为空")
    validate_file_signature(clean_name, content)
    if chunk_strategy not in {None, "fixed", "semantic"}:
        raise HTTPException(status_code=422, detail="无效的切块策略")

    digest = hashlib.sha256(content).hexdigest()
    document_id = str(uuid.uuid4())
    storage_path = settings.upload_dir / f"{document_id}{Path(clean_name).suffix.lower()}"
    now = utc_now()
    with db_session() as connection:
        kb_id = resolve_knowledge_base(
            connection, user, knowledge_base_id, "view" if is_seed else "upload"
        )
        groups, allowed_users = validate_document_acl(
            connection, user.org_id, visibility, group_ids, user_ids or []
        )
        duplicate = connection.execute(
            "SELECT id FROM documents WHERE org_id = ? AND sha256 = ?",
            (user.org_id, digest),
        ).fetchone()
        if duplicate:
            if is_seed:
                return str(duplicate["id"])
            raise HTTPException(status_code=409, detail="当前组织已经上传过相同内容")

        version_group_id = document_id
        version_number = 1
        if version_of:
            previous = connection.execute(
                "SELECT * FROM documents WHERE id = ? AND org_id = ? AND knowledge_base_id = ?",
                (version_of, user.org_id, kb_id),
            ).fetchone()
            if previous is None:
                raise HTTPException(status_code=404, detail="旧版本不存在或不属于当前知识库")
            require_knowledge_base_permission(connection, user, kb_id, "edit")
            version_group_id = previous["version_group_id"] or previous["id"]
            version_number = connection.execute(
                "SELECT COALESCE(MAX(version_number), 0) + 1 AS next_version FROM documents WHERE version_group_id = ?",
                (version_group_id,),
            ).fetchone()["next_version"]
            connection.execute(
                "UPDATE documents SET is_current = 0, archived_at = ? WHERE version_group_id = ?",
                (now, version_group_id),
            )

        effective_strategy = chunk_strategy or get_org_settings(user.org_id)["chunk_strategy"]
        connection.execute(
            """
            INSERT INTO documents
                (id, org_id, owner_id, knowledge_base_id, filename, title, content_type,
                 size_bytes, sha256, visibility, status, tags_json, version_group_id,
                 version_number, is_current, chunk_strategy, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'processing', ?, ?, ?, 1, ?, ?)
            """,
            (
                document_id,
                user.org_id,
                user.id,
                kb_id,
                clean_name,
                title.strip() or Path(clean_name).stem,
                content_type or "application/octet-stream",
                len(content),
                digest,
                visibility,
                json.dumps(list(dict.fromkeys(tags or [])), ensure_ascii=False),
                version_group_id,
                version_number,
                effective_strategy,
                now,
            ),
        )
        connection.executemany(
            "INSERT INTO document_groups (document_id, group_id) VALUES (?, ?)",
            [(document_id, group_id) for group_id in groups],
        )
        connection.executemany(
            "INSERT INTO document_users (document_id, user_id) VALUES (?, ?)",
            [(document_id, target_user_id) for target_user_id in allowed_users],
        )

    try:
        storage_path.write_bytes(content)
        if settings.task_queue == "celery" and not is_seed:
            from ..tasks import enqueue_document_index

            task_id = enqueue_document_index(document_id)
            with db_session() as connection:
                connection.execute(
                    "UPDATE documents SET task_id = ?, processing_stage = 'queued' WHERE id = ?",
                    (task_id, document_id),
                )
                write_audit(
                    connection,
                    user,
                    "document.upload.queued",
                    "document",
                    document_id,
                    metadata={"knowledge_base_id": kb_id, "version": version_number, "task_id": task_id},
                )
            return document_id
        chunk_count = await _index_document(document_id, content)
        if not is_seed:
            with db_session() as connection:
                write_audit(
                    connection,
                    user,
                    "document.upload",
                    "document",
                    document_id,
                    metadata={
                        "knowledge_base_id": kb_id,
                        "visibility": visibility,
                        "chunks": chunk_count,
                        "version": version_number,
                        "chunk_strategy": effective_strategy,
                    },
                )
    except HTTPException as exc:
        with db_session() as connection:
            connection.execute(
                "UPDATE documents SET status = 'failed', processing_stage = 'failed', error = ? WHERE id = ?",
                (str(exc.detail), document_id),
            )
        raise
    except Exception as exc:
        with db_session() as connection:
            connection.execute(
                "UPDATE documents SET status = 'failed', processing_stage = 'failed', error = ? WHERE id = ?",
                (str(exc)[:500], document_id),
            )
        raise HTTPException(status_code=503, detail="文档索引服务暂时不可用，请稍后重试") from exc
    return document_id


async def reparse_document(document_id: str, user: CurrentUser) -> int:
    with db_session() as connection:
        document = connection.execute(
            "SELECT * FROM documents WHERE id = ? AND org_id = ?",
            (document_id, user.org_id),
        ).fetchone()
        if document is None:
            raise HTTPException(status_code=404, detail="文档不存在")
        require_knowledge_base_permission(
            connection, user, document["knowledge_base_id"], "edit"
        )
        connection.execute(
            "UPDATE documents SET status = 'processing', processing_stage = 'queued', error = NULL WHERE id = ?",
            (document_id,),
        )
    path = get_settings().upload_dir / f"{document_id}{Path(document['filename']).suffix.lower()}"
    if not path.exists():
        raise HTTPException(status_code=404, detail="原始文件不存在，无法重新解析")
    try:
        count = await _index_document(document_id, path.read_bytes())
    except Exception:
        with db_session() as connection:
            connection.execute(
                "UPDATE documents SET status = 'failed', error = '重新解析失败' WHERE id = ?",
                (document_id,),
            )
        raise
    with db_session() as connection:
        write_audit(
            connection,
            user,
            "document.reparse",
            "document",
            document_id,
            metadata={"chunks": count},
        )
    return count
